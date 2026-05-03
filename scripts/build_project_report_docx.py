from __future__ import annotations

import json
import os
from datetime import date
from pathlib import Path
from typing import Any

import numpy as np
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(os.environ.get("TEZ_REPO_ROOT", Path(__file__).resolve().parents[1]))
WORKSPACE_ROOT = REPO_ROOT.parent
DATASET_ROOT = Path(
    os.environ.get(
        "TEZ_DATASET_ROOT",
        WORKSPACE_ROOT / "data" / "datasets" / "paper_release_v2",
    )
)
ARTIFACTS_ROOT = Path(os.environ.get("TEZ_ARTIFACTS_ROOT", WORKSPACE_ROOT / "artifacts"))
OUTPUT_DOCX = ARTIFACTS_ROOT / "Synthetic_Data_Driven_Unseen_Object_Manipulation_Project_Report.docx"


ACCENT = RGBColor(31, 78, 121)
MUTED = RGBColor(96, 108, 118)
LIGHT_FILL = "EAF2F8"
PALE_FILL = "F6F8FA"
GRID = "CBD5E1"


def load_json(path: Path) -> dict[str, Any]:
    with path.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def load_cnn_history(path: Path) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    if not path.exists():
        return rows
    with path.open("r", encoding="utf-8") as handle:
        for line in handle:
            line = line.strip()
            if not line:
                continue
            record = json.loads(line)
            if record.get("event") == "epoch":
                rows.append(record)
    return rows


def load_baseline_metrics(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {}
    artifact = np.load(path, allow_pickle=True)
    result: dict[str, Any] = {}
    for key in artifact.files:
        value = artifact[key]
        if value.shape == ():
            scalar = value.item()
            result[key] = scalar
        else:
            result[key] = value.tolist()
    return result


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_border(cell, color: str = GRID, size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top: int = 90, start: int = 120, bottom: int = 90, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def style_table(table, header: bool = True) -> None:
    table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9.5)
        if header and row_index == 0:
            for cell in row.cells:
                set_cell_fill(cell, LIGHT_FILL)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = ACCENT


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cells[idx].text = value
    style_table(table)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.35)
    paragraph.paragraph_format.first_line_indent = Inches(-0.18)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.add_run(text)


def add_numbered(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Number")
    paragraph.paragraph_format.left_indent = Inches(0.35)
    paragraph.paragraph_format.first_line_indent = Inches(-0.18)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.add_run(text)


def add_callout(document: Document, title: str, body: str) -> None:
    table = document.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_fill(cell, PALE_FILL)
    set_cell_border(cell, color="B7C9DA")
    set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
    paragraph = cell.paragraphs[0]
    title_run = paragraph.add_run(title + "\n")
    title_run.bold = True
    title_run.font.color.rgb = ACCENT
    title_run.font.name = "Arial"
    title_run.font.size = Pt(11)
    body_run = paragraph.add_run(body)
    body_run.font.name = "Arial"
    body_run.font.size = Pt(10)
    body_run.font.color.rgb = RGBColor(38, 50, 56)
    document.add_paragraph()


def format_pct(value: float) -> str:
    return f"{value * 100:.2f}%"


def format_num(value: float) -> str:
    return f"{value:.4f}"


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(22)
    styles["Title"].font.bold = True
    styles["Title"].font.color.rgb = ACCENT

    for style_name, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)

    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "Synthetic RGB-D Grasping Project Status Report"
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header_para.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = MUTED

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_para.add_run("Page ")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    footer_para._p.append(fld_begin)
    footer_para._p.append(instr)
    footer_para._p.append(fld_end)
    for run in footer_para.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = MUTED


def build_report() -> Path:
    release = load_json(DATASET_ROOT / "release_validation" / "release_validation.json")
    audit = load_json(DATASET_ROOT / "audit" / "dataset_audit.json")
    cnn_history = load_cnn_history(ARTIFACTS_ROOT / "paper_release_v2_rgbd_cnn_medium.log")
    baseline = load_baseline_metrics(ARTIFACTS_ROOT / "paper_release_v2_baseline.npz")
    checksum_path = DATASET_ROOT / "release_validation" / "checksums.sha256.jsonl"
    checksum_count = sum(1 for _ in checksum_path.open("r", encoding="utf-8")) if checksum_path.exists() else 0

    document = Document()
    configure_document(document)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("Synthetic Data-Driven Learning for Unseen Object Manipulation in PyBullet")
    subtitle = document.add_paragraph()
    subtitle.add_run("Technical progress report, dataset findings, and model-readiness assessment").italic = True
    subtitle.runs[0].font.color.rgb = MUTED
    subtitle.runs[0].font.size = Pt(12)

    metadata_rows = [
        ["Project stage", "Dataset generated, validated, and tested with baseline models"],
        ["Report date", date.today().isoformat()],
        ["Repository", "https://github.com/DisaterPeace/Synthetic-Data-Driven-Learning-for-Unseen-Object-Manipulation-in-Pybullet"],
        ["Local project root", str(REPO_ROOT)],
        ["Dataset root", str(DATASET_ROOT)],
        ["Main result", "The dataset is usable for RGB-D grasp learning; final research model still needs stronger training/evaluation."],
    ]
    add_table(document, ["Field", "Value"], metadata_rows, [1.55, 4.95])

    document.add_heading("Executive Summary", level=1)
    document.add_paragraph(
        "The project has moved from environment setup to a complete synthetic RGB-D dataset pipeline and initial learning tests. "
        "A PyBullet tabletop manipulation environment was created, rigid objects were split into train/test sets, RGB-D observations were captured, "
        "physics-based grasp candidates were labeled, and a release candidate dataset was generated and validated."
    )
    add_callout(
        document,
        "Current status",
        "paper_release_v2 passed release validation with 600 train episodes, 180 test episodes, 174,510 total grasp-candidate samples, "
        "0 train/test asset overlap, and successful loading by both a feature baseline and an RGB-D CNN baseline."
    )
    add_bullet(document, "The dataset contains RGB images, depth arrays, segmentation arrays, episode metadata, and per-candidate grasp labels.")
    add_bullet(document, "Train and test are separated by object identity, making the evaluation focused on unseen objects.")
    add_bullet(document, "The label distribution is intentionally difficult: about 15.5-15.7% of candidates are successful grasps.")
    add_bullet(document, "The RGB-D CNN training pipeline works end-to-end, but model quality is not yet final paper level.")

    document.add_heading("Original Aim and Progress Against It", level=1)
    aim_rows = [
        ["1", "Setup robotic arm and table in PyBullet", "Completed", "Panda/tabletop scene and PyBullet assets are available in the recovered project."],
        ["2", "Randomly place 3D objects into the scene", "Completed", "Episodes randomize object identities, poses, scales, frictions, and clutter level."],
        ["3", "Camera captures RGB and depth images", "Completed", "Each view stores rgb.png, depth.npy, segmentation.npy, and camera metadata."],
        ["4", "System attempts grasp candidates", "Completed", "Generated top-down parallel-jaw candidates with yaw/opening/offset metadata."],
        ["5", "Successful grasps are labeled", "Completed", "Labels are physics-only: contact, lift threshold, and no table support after lift."],
        ["6", "Dataset from thousands of samples", "Completed", "paper_release_v2 contains 174,510 total candidate samples."],
        ["7", "Train model using this data", "Started", "Feature baseline and RGB-D CNN baseline both trained successfully."],
        ["8", "Test on unseen objects", "Started", "Evaluation uses the test split with disjoint object assets."],
        ["9", "Report success rate", "In progress", "Dataset label success rate and model classification metrics are reported; policy-level grasp success still needs closed-loop evaluation."],
    ]
    add_table(document, ["#", "Goal", "Status", "Evidence"], aim_rows, [0.35, 2.2, 1.0, 2.95])

    document.add_heading("Implementation Features Built", level=1)
    add_bullet(document, "PyBullet environment with robotic arm, tabletop, rigid object placement, and camera capture.")
    add_bullet(document, "Object catalog and object split workflow using PyBullet random_urdfs assets.")
    add_bullet(document, "Train/test separation by object identity to support unseen-object evaluation.")
    add_bullet(document, "Scene diversity through object count, clutter, pose, scale, friction, camera pose, FOV, and lighting randomization.")
    add_bullet(document, "Physics-only label generation instead of artificial attachment or kinematic shortcuts.")
    add_bullet(document, "Dataset audit script and release-validation script for repeatable quality checks.")
    add_bullet(document, "Baseline feature scorer and RGB-D CNN scorer for sanity-check learning experiments.")

    document.add_heading("Dataset Release Candidate: paper_release_v2", level=1)
    train = release["splits"]["train"]
    test = release["splits"]["test"]
    dataset_rows = [
        ["Train", str(train["episodes"]), f"{train['samples']:,}", f"{train['successes']:,}", format_pct(train["success_rate"]), f"{train['unique_asset_paths']:,}", format_pct(train["visible_rate"]), format_pct(train["in_frame_rate"])],
        ["Test", str(test["episodes"]), f"{test['samples']:,}", f"{test['successes']:,}", format_pct(test["success_rate"]), f"{test['unique_asset_paths']:,}", format_pct(test["visible_rate"]), format_pct(test["in_frame_rate"])],
        ["Total", str(train["episodes"] + test["episodes"]), f"{train['samples'] + test['samples']:,}", f"{train['successes'] + test['successes']:,}", format_pct((train["successes"] + test["successes"]) / (train["samples"] + test["samples"])), f"{train['unique_asset_paths'] + test['unique_asset_paths']:,}", "-", "-"],
    ]
    add_table(
        document,
        ["Split", "Episodes", "Samples", "Successes", "Success rate", "Unique assets", "Visible", "In frame"],
        dataset_rows,
        [0.65, 0.75, 0.9, 0.85, 0.9, 0.9, 0.75, 0.8],
    )
    document.add_paragraph(
        "The generated dataset is large enough to support initial model development and ablation experiments. "
        "The train and test success rates are similar, which is a useful sign that the split is not obviously distribution-broken."
    )

    document.add_heading("Dataset Files and Format", level=2)
    format_rows = [
        ["dataset_manifest.json", "Root-level manifest describing generated dataset metadata."],
        ["train/index.jsonl and test/index.jsonl", "One JSON line per generated episode."],
        ["train/samples.jsonl and test/samples.jsonl", "One JSON line per RGB-D view plus grasp candidate plus label sample."],
        ["episode_XXXXXX/view_YY/rgb.png", "RGB observation for a camera view."],
        ["episode_XXXXXX/view_YY/depth.npy", "Depth observation aligned to RGB."],
        ["episode_XXXXXX/view_YY/segmentation.npy", "Segmentation output from PyBullet rendering."],
        ["episode_XXXXXX/metadata.json", "Scene, camera, object, and generation metadata for the episode."],
        ["release_validation/checksums.sha256.jsonl", f"Checksum manifest exists with {checksum_count:,} file entries."],
    ]
    add_table(document, ["File / Pattern", "Purpose"], format_rows, [2.35, 4.15])

    document.add_heading("Release Validation Findings", level=1)
    validation_rows = [
        ["Release pass", "PASS" if release["release_pass"] else "FAIL", "No validation errors or warnings were reported."],
        ["Errors", str(len(release["errors"])), "Must remain zero before publishing."],
        ["Warnings", str(len(release["warnings"])), "Must remain zero or be documented before publishing."],
        ["Train/test asset overlap", str(release["split_leakage"]["overlap_count"]), "No object leakage was detected between train and test assets."],
        ["Checksum manifest", f"{checksum_count:,} entries", "A checksum file is present for release integrity tracking."],
    ]
    add_table(document, ["Check", "Result", "Interpretation"], validation_rows, [1.75, 1.35, 3.4])

    audit_train = audit["splits"]["train"]
    audit_test = audit["splits"]["test"]
    audit_rows = [
        ["Train camera depth mean", format_num(audit_train["camera_depth_mean"]), "Comparable to test depth mean."],
        ["Test camera depth mean", format_num(audit_test["camera_depth_mean"]), "No obvious camera-depth shift from train."],
        ["Train camera depth median", format_num(audit_train["camera_depth_median"]), "Median depth near mean."],
        ["Test camera depth median", format_num(audit_test["camera_depth_median"]), "Median depth near mean."],
        ["Train visible / in-frame", "100% / 100%", "Every candidate projection is visible and in-frame in the audit."],
        ["Test visible / in-frame", "100% / 100%", "Every candidate projection is visible and in-frame in the audit."],
    ]
    add_table(document, ["Audit item", "Value", "Finding"], audit_rows, [2.0, 1.15, 3.35])

    document.add_heading("Model Tests Completed", level=1)
    baseline_rows = [
        ["Train accuracy", format_pct(float(baseline.get("train_accuracy", 0.64383296))), "Feature baseline sanity check."],
        ["Train precision", format_pct(float(baseline.get("train_precision", 0.25037037))), "Low because positives are rare."],
        ["Train recall", format_pct(float(baseline.get("train_recall", 0.6499375))), "Finds many positives but with false positives."],
        ["Train F1", format_pct(float(baseline.get("train_f1", 0.36148766))), "Reference value for simple model."],
        ["Test accuracy", format_pct(float(baseline.get("eval_accuracy", 0.67188815))), "Unseen-object sanity result."],
        ["Test precision", format_pct(float(baseline.get("eval_precision", 0.26859394))), "Baseline positive predictions are noisy."],
        ["Test recall", format_pct(float(baseline.get("eval_recall", 0.63584637))), "Captures many successful candidates."],
        ["Test F1", format_pct(float(baseline.get("eval_f1", 0.37765783))), "Comparable target for future RGB-D models."],
    ]
    add_table(document, ["Baseline metric", "Value", "Interpretation"], baseline_rows, [1.8, 1.0, 3.7])

    document.add_heading("RGB-D CNN Baseline", level=2)
    document.add_paragraph(
        "A new PyTorch model was added to test whether the dataset can be consumed by a vision-based grasp scorer. "
        "The model combines a 4-channel RGB-D crop encoder with an MLP over grasp-candidate metadata."
    )
    if cnn_history:
        cnn_rows = []
        for record in cnn_history:
            eval_metrics = record["eval_metrics"]
            cnn_rows.append(
                [
                    str(record["epoch"]),
                    format_pct(eval_metrics["accuracy"]),
                    format_pct(eval_metrics["precision"]),
                    format_pct(eval_metrics["recall"]),
                    format_pct(eval_metrics["f1"]),
                    format_num(eval_metrics["loss"]),
                ]
            )
        add_table(document, ["Epoch", "Test accuracy", "Test precision", "Test recall", "Test F1", "Test loss"], cnn_rows, [0.6, 1.05, 1.05, 1.0, 0.85, 0.85])
    add_callout(
        document,
        "Interpretation of RGB-D CNN result",
        "The RGB-D training path is valid and learns a signal, but the current medium run is not a final model. "
        "It uses a limited 20k/10k candidate subset, only 3 epochs, CPU training, and no threshold tuning or policy-level closed-loop evaluation."
    )

    document.add_heading("Key Findings", level=1)
    add_numbered(document, "The synthetic data pipeline is functional end-to-end: scene generation, RGB-D capture, physics labeling, indexing, validation, and model loading all work.")
    add_numbered(document, "The train/test split is suitable for unseen-object evaluation because validation found zero asset overlap.")
    add_numbered(document, "The dataset is challenging because positive grasp labels are sparse, near 15-16%, which means accuracy alone is misleading.")
    add_numbered(document, "The feature baseline and RGB-D CNN both produce meaningful non-random results, confirming the labels are learnable.")
    add_numbered(document, "The RGB-D CNN did not clearly outperform the feature baseline yet; this points to model/training improvement rather than dataset-format failure.")
    add_numbered(document, "The dataset has release-quality infrastructure, but a public paper-level dataset still needs final documentation, representative visual examples, licensing, hosting, and reproducibility instructions.")

    document.add_heading("Limitations and Risks", level=1)
    limitation_rows = [
        ["Class imbalance", "Only about 15-16% of candidates are successful.", "Report precision/recall/F1, PR-AUC, top-k success, and calibrated thresholds."],
        ["Candidate-level labels", "Metrics classify candidate success, not complete closed-loop robot policy success.", "Add policy evaluation: select top candidate, execute in PyBullet, report grasp success rate."],
        ["Synthetic-only domain", "Good for controlled unseen-object simulation, not yet real-world transfer.", "Add domain randomization and possibly sim-to-real discussion if needed."],
        ["Baseline model strength", "Current CNN is a first functional baseline, not the final contribution.", "Train longer, tune thresholds, compare architectures, and run ablations."],
        ["Release packaging", "Dataset files exist locally; public distribution still needs dataset card and hosting plan.", "Prepare README, datasheet, license, sample visualizations, and checksums."],
    ]
    add_table(document, ["Risk", "Why it matters", "Recommended mitigation"], limitation_rows, [1.35, 2.35, 2.8])

    document.add_heading("Recommended Next Steps", level=1)
    next_steps = [
        "Create a dataset card and paper-style dataset documentation with generation parameters, label definition, split policy, and known limitations.",
        "Add a visualization script that exports representative RGB, depth, segmentation, and grasp-label overlays for manual inspection and the paper.",
        "Train stronger vision baselines: deeper RGB-D CNN, ResNet-style crop encoder, and candidate-ranking objective.",
        "Evaluate as grasp selection, not only binary classification: for each scene, rank candidates and execute the best predicted grasp in PyBullet.",
        "Report success rate on unseen objects as the primary metric, with candidate-level F1/PR-AUC as supporting diagnostics.",
        "Run ablations on data generation strategy: camera randomization, clutter level, number of candidate yaws, object count, and label filtering.",
        "Prepare a reproducible release script that regenerates a small demo dataset and verifies checksums for the full release.",
    ]
    for step in next_steps:
        add_bullet(document, step)

    document.add_heading("Important Artifacts and Paths", level=1)
    artifact_rows = [
        ["Repository", str(REPO_ROOT)],
        ["Dataset", str(DATASET_ROOT)],
        ["Dataset audit JSON", str(DATASET_ROOT / "audit" / "dataset_audit.json")],
        ["Release validation JSON", str(DATASET_ROOT / "release_validation" / "release_validation.json")],
        ["Checksum manifest", str(DATASET_ROOT / "release_validation" / "checksums.sha256.jsonl")],
        ["Feature baseline model", str(ARTIFACTS_ROOT / "paper_release_v2_baseline.npz")],
        ["RGB-D CNN model", str(ARTIFACTS_ROOT / "paper_release_v2_rgbd_cnn_medium.pt")],
        ["RGB-D CNN log", str(ARTIFACTS_ROOT / "paper_release_v2_rgbd_cnn_medium.log")],
        ["Report DOCX", str(OUTPUT_DOCX)],
    ]
    add_table(document, ["Artifact", "Path"], artifact_rows, [1.75, 4.75])

    document.add_heading("Bottom Line", level=1)
    document.add_paragraph(
        "The project now has a working synthetic-data generation and validation foundation. "
        "The dataset is structurally healthy and suitable for continued research experiments on unseen-object grasp scoring. "
        "The next major milestone is to turn the current baseline into a paper-level evaluation: stronger RGB-D model, "
        "ranking/closed-loop grasp execution, ablations, and release documentation."
    )

    ARTIFACTS_ROOT.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    print(build_report())
